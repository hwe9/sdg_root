import os
import csv
from docx import Document
import PyPDF2
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import datetime

class FileHandler:
    def __init__(self, images_dir):
        self.images_dir = images_dir
        os.makedirs(self.images_dir, exist_ok=True)

    def _extract_images_from_pdf(self, file_path: str, doc_name: str) -> list:
        """Extrahiert Bilder aus einer PDF-Datei und speichert sie."""
        image_paths = []
        try:
            pdf_document = fitz.open(file_path)
            for page_index in range(len(pdf_document)):
                page = pdf_document[page_index]
                image_list = page.get_images(full=True)
                for image_index, img in enumerate(image_list, start=1):
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    file_extension = base_image["ext"]
                    image_path = os.path.join(self.images_dir, f"{doc_name}_page{page_index+1}_img{image_index}.{file_extension}")
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                    image_paths.append({
                        "original_path": image_path,
                        "page": page_index + 1,
                        "caption": None,  # Optional bei OCR/Meta extrahieren!
                        "sdg_tags": {},
                        "ai_tags": None,
                        "image_type": file_extension
                    })
        except Exception as e:
            print(f"Fehler beim Extrahieren von Bildern aus PDF: {e}")
        return image_paths

    def _extract_images_from_docx(self, file_path: str, doc_name: str) -> list:
        """Extrahiert Bilder aus einer DOCX-Datei und speichert sie."""
        image_paths = []
        try:
            document = Document(file_path)
            for rel in document.part.rels:
                if "image" in document.part.rels[rel].target_ref:
                    image_part = document.part.rels[rel].target_part
                    image_path = os.path.join(self.images_dir, f"{doc_name}_{os.path.basename(image_part.partname)}")
                    with open(image_path, "wb") as f:
                        f.write(image_part.blob)
                    image_paths.append({
                        "original_path": image_path,
                        "page": None,
                        "caption": None,
                        "sdg_tags": {},
                        "ai_tags": None,
                        "image_type": os.path.splitext(image_path)[-1].replace('.', '')
                    })
        except Exception as e:
            print(f"Fehler beim Extrahieren von Bildern aus DOCX: {e}")
        return image_paths

    def get_text_from_file(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def extract_text_from_pdf(self, file_path: str) -> tuple:
        text = ""
        metadata = {}
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                doc_info = pdf_reader.metadata
                if doc_info:
                    metadata['title'] = doc_info.get('/Title', None)
                    metadata['authors'] = doc_info.get('/Author', None)
                    metadata['creation_date'] = doc_info.get('/CreationDate', None)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except Exception as e:
            print(f"Fehler beim Extrahieren von Text oder Metadaten aus PDF: {e}")
        return text, metadata
    
    def convert_pdf_to_txt(self, pdf_path: str, txt_path: str = None) -> str:
        """Extract text from PDF and save as .txt file.
        
        Args:
            pdf_path: Path to input PDF file
            txt_path: Optional output path. If None, uses same path with .txt extension
            
        Returns:
            Path to created .txt file
        """
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        if not pdf_path.lower().endswith('.pdf'):
            raise ValueError(f"File is not a PDF: {pdf_path}")
        
        # Extract text
        text, metadata = self.extract_text_from_pdf(pdf_path)
        
        if not text.strip():
            print(f"Warning: No text extracted from PDF: {pdf_path}")
        
        # Determine output path
        if txt_path is None:
            base_path = os.path.splitext(pdf_path)[0]
            txt_path = f"{base_path}.txt"
        
        # Ensure output directory exists
        output_dir = os.path.dirname(txt_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        
        # Save as .txt file
        try:
            with open(txt_path, 'w', encoding='utf-8') as f:
                # Optional: Add metadata header
                if metadata.get('title'):
                    f.write(f"Title: {metadata['title']}\n")
                if metadata.get('authors'):
                    f.write(f"Authors: {metadata['authors']}\n")
                if metadata.get('creation_date'):
                    f.write(f"Creation Date: {metadata['creation_date']}\n")
                if any(metadata.values()):
                    f.write("\n" + "="*80 + "\n\n")
                
                f.write(text)
            
            print(f"✅ Converted PDF to TXT: {pdf_path} -> {txt_path}")
            return txt_path
        except Exception as e:
            raise IOError(f"Failed to save TXT file: {e}")

    def extract_text_from_docx(self, file_path: str) -> str:
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Fehler beim Extrahieren von Text aus DOCX: {e}")
        return text

    def extract_text_from_csv(self, file_path: str) -> str:
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                csv_reader = csv.reader(f)
                for row in csv_reader:
                    text += " ".join(row) + "\n"
        except Exception as e:
            print(f"Fehler beim Extrahieren von Text aus CSV: {e}")
        return text

    def extract_text_from_html(self, file_path: str) -> str:
        """Extract text from HTML file, removing HTML tags."""
        text = ""
        try:
            import re
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
                # Remove script and style elements
                html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                # Remove HTML tags
                text = re.sub(r'<[^>]+>', '', html_content)
                # Clean up whitespace
                text = re.sub(r'\s+', ' ', text)
                text = text.strip()
        except Exception as e:
            print(f"Fehler beim Extrahieren von Text aus HTML: {e}")
        return text

    def extract_text(self, file_path: str) -> str:
        if file_path.endswith('.mp3'):
            return ""
        elif file_path.endswith('.pdf'):
            text, _ = self.extract_text_from_pdf(file_path)
            return text
        elif file_path.endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        elif file_path.endswith('.csv'):
            return self.extract_text_from_csv(file_path)
        elif file_path.endswith('.html') or file_path.endswith('.htm'):
            return self.extract_text_from_html(file_path)
        else:
            return self.get_text_from_file(file_path)

    def get_metadata_from_json(self, file_path: str) -> dict:
        import json
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            return {}

    def cleanup_processed_data(self, directory: str, retention_days: int):
        """Löscht Dateien in einem Verzeichnis, die älter als retention_days sind."""
        print(f"Starte Bereinigung von {directory}...")
        now = datetime.datetime.now()
        for filename in os.listdir(directory):
            filepath = os.path.join(directory, filename)
            if os.path.isfile(filepath):
                file_creation_time = datetime.datetime.fromtimestamp(os.path.getctime(filepath))
                if (now - file_creation_time).days > retention_days:
                    print(f"Lösche alte Datei: {filepath}")
                    os.remove(filepath)
        print("Bereinigung abgeschlossen.")
